import streamlit as st
import os
from docx import Document
from docx2pdf import convert
import tkinter as tk
from tkinter import filedialog

def seleccionar_carpeta(titulo="Selecciona una carpeta"):
    root = tk.Tk()
    root.withdraw()  # Oculta la ventana principal de Tkinter
    carpeta = filedialog.askdirectory(title=titulo)
    root.destroy()
    return carpeta

def reemplazar_en_documento(ruta_entrada, ruta_salida, reemplazos):
    doc = Document(ruta_entrada)

    # Reemplazar en párrafos
    for p in doc.paragraphs:
        for buscar, reemplazar in reemplazos.items():
            if buscar in p.text:
                p.text = p.text.replace(buscar, reemplazar)

    # Reemplazar en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for buscar, reemplazar in reemplazos.items():
                    if buscar in celda.text:
                        celda.text = celda.text.replace(buscar, reemplazar)

    doc.save(ruta_salida)

# --- Streamlit UI ---
st.title("🔄 Reemplazo Masivo en Word y Exportación a PDF")

if st.button("📁 Seleccionar carpeta de entrada"):
    carpeta = seleccionar_carpeta("Selecciona la carpeta con archivos .docx")
    st.session_state['carpeta'] = carpeta
    st.success(f"Carpeta seleccionada: {carpeta}")

if st.button("📂 Seleccionar carpeta de salida"):
    carpeta_salida = seleccionar_carpeta("Selecciona la carpeta de salida para los PDFs")
    st.session_state['carpeta_salida'] = carpeta_salida
    st.success(f"Carpeta de salida: {carpeta_salida}")

# Diccionario de reemplazos
st.markdown("✏️ **Agrega pares de texto a buscar y reemplazar**")
reemplazos = {}
num_pares = st.number_input("Número de pares búsqueda/reemplazo", min_value=1, max_value=20, value=1, step=1)

for i in range(num_pares):
    buscar = st.text_input(f"🔎 Buscar texto #{i+1}", key=f"buscar_{i}")
    reemplazar = st.text_input(f"✏️ Reemplazar por #{i+1}", key=f"reemplazar_{i}")
    if buscar and reemplazar:
        reemplazos[buscar] = reemplazar

# Procesar botón
if st.button("🚀 Procesar documentos y exportar a PDF"):
    carpeta = st.session_state.get('carpeta')
    carpeta_salida = st.session_state.get('carpeta_salida')

    if not carpeta or not os.path.exists(carpeta):
        st.error("❌ Carpeta de entrada no válida")
    elif not carpeta_salida or not os.path.exists(carpeta_salida):
        st.error("❌ Carpeta de salida no válida")
    else:
        archivos = [f for f in os.listdir(carpeta) if f.endswith(".docx")]
        total = len(archivos)
        if total == 0:
            st.warning("⚠️ No hay archivos .docx en la carpeta seleccionada.")
        else:
            st.info(f"🔄 Procesando {total} documentos...")
            for archivo in archivos:
                ruta_docx = os.path.join(carpeta, archivo)
                nombre_modificado = f"MOD_{archivo}"
                ruta_modificado = os.path.join(carpeta_salida, nombre_modificado)
                
                # Reemplazar texto y guardar el nuevo .docx
                reemplazar_en_documento(ruta_docx, ruta_modificado, reemplazos)

                # Convertir a PDF
                ruta_pdf = ruta_modificado.replace(".docx", ".pdf")
                convert(ruta_modificado, ruta_pdf)

                st.success(f"✅ PDF generado: {os.path.basename(ruta_pdf)}")
            
            st.balloons()
            st.info(f"🎉 ¡Todos los documentos ({total}) han sido procesados y convertidos a PDF!")



